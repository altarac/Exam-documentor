import pandas as pd
from docx import Document
from docx.shared import Inches

questionsdb = pd.read_excel('Elem Exam Question Makerspace (Responses).xlsx')
i=0

document = Document()

document.add_heading('Documentation Grade 4 Exam Questions', 0)



for index, row in questionsdb.iterrows():
    i = i+1
    document.add_heading('Question ' + str(i) + '\n', 1)
    p = document.add_paragraph('')
    p.add_run('Targetted progression of learning: \n')
    p.add_run(str(row[0]) + '\n')
    p.add_run('Question context: \n')
    p.add_run(str(row[1]) + '\n')       
    p.add_run('Question posed: \n')
    p.add_run(row[4] + '\n')

    p.add_run('Option A: ')
    p.add_run(row[5])
    p.add_run('Option B: ')
    p.add_run(row[6])
    p.add_run('Option C: ')
    p.add_run(row[7])
    p.add_run('Option D: ')
    p.add_run(row[8] + '\n')

    p.add_run('Correct solution: ')
    p.add_run(row[9] + '\n')
    p.add_run('Image:' + '\n')
    if pd.isnull(row[2]) and pd.isnull(row[3]):
        p.add_run('NONE \n')
    else:
        #p.add_run(row[2] + '\n')
        document.add_picture('imgB1.png', width=Inches(4.0))
    document.add_page_break()
     
document.save('Documentation Grade 4 Exam Questions.docx')





#p = document.add_paragraph('A plain paragraph having some ')
#p.add_run('bold').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True

#document.add_heading('Heading, level 1', level=1)
#document.add_paragraph('Intense quote', style='IntenseQuote')

#document.add_paragraph(
#    'first item in unordered list', style='ListBullet'
#)
#document.add_paragraph(
#    'first item in ordered list', style='ListNumber'
#)

#document.add_picture('monty-truth.png', width=Inches(1.25))

#table = document.add_table(rows=1, cols=3)
#hdr_cells = table.rows[0].cells
#hdr_cells[0].text = 'Qty'
#hdr_cells[1].text = 'Id'
#hdr_cells[2].text = 'Desc'


#document.add_page_break()

#document.save('demo.docx')

